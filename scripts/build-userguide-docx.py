"""
Builds Yardward-Pro-User-Guide.docx with deeply-detailed content + 48 screenshots.

Output: docs/Yardward-Pro-User-Guide.docx
Then converts to PDF via docx2pdf (uses MS Word COM automation on Windows).
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent.parent
SHOTS = ROOT / "docs" / "screenshots"
DOCX = ROOT / "docs" / "Yardward-Pro-User-Guide.docx"
PDF = ROOT / "docs" / "Yardward-Pro-User-Guide.pdf"

AMBER = RGBColor(0xC2, 0x6A, 0x00)
NAVY = RGBColor(0x0E, 0x1E, 0x3D)
MUTED = RGBColor(0x6B, 0x7A, 0x90)
SUCCESS = RGBColor(0x0F, 0x7A, 0x3D)
DANGER = RGBColor(0xB7, 0x29, 0x29)
LIGHT_GRAY = RGBColor(0xEE, 0xEE, 0xEE)


# ============================================================================
# Helper functions
# ============================================================================

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if color:
            run.font.color.rgb = color
    return h


def add_para(doc, text, size=11, bold=False, color=None, italic=False, indent=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def add_steps(doc, steps):
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(style='List Number')
        r = p.add_run(s)
        r.font.size = Pt(11)


def add_bullets(doc, items):
    for s in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(s)
        r.font.size = Pt(11)


def add_callout(doc, kind, body):
    """kind: 'tip' | 'warning' | 'info'"""
    color = {'tip': SUCCESS, 'warning': DANGER, 'info': NAVY}[kind]
    label = {'tip': '💡 TIP', 'warning': '⚠️ WARNING', 'info': 'ℹ️ NOTE'}[kind]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.color.rgb = color
    r1.font.size = Pt(10)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = color


def add_field_table(doc, fields):
    """fields: list of (Field, Description) tuples"""
    table = doc.add_table(rows=len(fields) + 1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.autofit = True
    hdr = table.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Description'
    for c in hdr:
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for i, (f, d) in enumerate(fields, 1):
        row = table.rows[i].cells
        row[0].text = f
        row[1].text = d
        for c in row:
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(10)


def add_image(doc, name, caption=None, width_in=6.0):
    path = SHOTS / name
    if not path.exists():
        add_para(doc, f"[missing screenshot: {name}]", italic=True, color=MUTED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = MUTED


def add_toc_entry(doc, label, page_hint=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.font.size = Pt(11)


def add_section_break(doc):
    doc.add_page_break()


# ============================================================================
# Document
# ============================================================================

def build():
    doc = Document()

    # Set base style
    for s in doc.styles:
        try:
            if s.name == "Normal":
                s.font.name = "Calibri"
                s.font.size = Pt(11)
        except Exception:
            pass
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ------------------------------------------------------------------------
    # COVER
    # ------------------------------------------------------------------------
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Yardward Pro")
    r.font.size = Pt(48)
    r.bold = True
    r.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Trucking & Haulage Operations CRM")
    sr.font.size = Pt(18)
    sr.font.italic = True
    sr.font.color.rgb = AMBER

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr2 = sub2.add_run("Complete User Guide")
    sr2.font.size = Pt(20)
    sr2.bold = True
    sr2.font.color.rgb = MUTED

    for _ in range(8):
        doc.add_paragraph()

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("Version 1.0 · Edition: Production · 2026")
    fr.font.size = Pt(11)
    fr.font.color.rgb = MUTED

    foot2 = doc.add_paragraph()
    foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr2 = foot2.add_run("yardward-pro.vercel.app")
    fr2.font.size = Pt(11)
    fr2.font.color.rgb = AMBER

    # ------------------------------------------------------------------------
    # TABLE OF CONTENTS
    # ------------------------------------------------------------------------
    add_section_break(doc)
    add_heading(doc, "Table of Contents", level=1, color=AMBER)
    toc_entries = [
        "1. Introduction",
        "2. Getting started",
        "    2.1. Sign-in walkthrough",
        "    2.2. First-time sign-in & password rotation",
        "    2.3. Forgot password recovery",
        "    2.4. Installing the mobile PWA",
        "    2.5. Roles & access overview",
        "3. Admin guide",
        "    3.1. Dashboard",
        "    3.2. Schedule",
        "    3.3. Jobs",
        "    3.4. Drivers & mechanics",
        "    3.5. Vehicles",
        "    3.6. Live vehicle map",
        "    3.7. Work orders",
        "    3.8. Communications",
        "    3.9. Timesheets",
        "    3.10. SMS log",
        "    3.11. Purchase orders",
        "    3.12. Prepaid tickets",
        "    3.13. Clients",
        "    3.14. Forms & submissions",
        "    3.15. Tenders",
        "    3.16. Error log",
        "    3.17. Reports",
        "    3.18. Settings",
        "4. Driver guide (mobile)",
        "    4.1. Home",
        "    4.2. My jobs",
        "    4.3. Forms tile menu",
        "    4.4. Start of day & clock-in",
        "    4.5. Tool checklist",
        "    4.6. Vehicle inspection (pre-trip)",
        "    4.7. Job log",
        "    4.8. Work order submission",
        "    4.9. End of day & clock-out",
        "    4.10. Prepaid tickets",
        "    4.11. Messages",
        "    4.12. Profile",
        "5. Mechanic guide",
        "    5.1. Workshop dashboard",
        "    5.2. Work orders queue",
        "    5.3. Messages",
        "    5.4. Purchase requests",
        "    5.5. Maintenance logs",
        "    5.6. Parts inventory",
        "6. Tokenized driver links (QR access)",
        "7. Offline handling",
        "8. Notifications & alerts",
        "9. Integrations",
        "    9.1. Twilio (SMS & Communications)",
        "    9.2. Geotab (GPS & telematics)",
        "    9.3. QuickBooks Online (invoicing & payroll)",
        "    9.4. Fleetio (vehicle data import)",
        "    9.5. Resend (tender digest emails)",
        "10. Security & data isolation",
        "11. Support & troubleshooting",
        "12. Glossary",
    ]
    for e in toc_entries:
        add_toc_entry(doc, e)

    # ========================================================================
    # SECTION 1 — INTRODUCTION
    # ========================================================================
    add_section_break(doc)
    add_heading(doc, "1. Introduction", level=1, color=AMBER)

    add_para(doc, "Yardward Pro is a single, integrated operations platform for trucking and haulage companies. It consolidates scheduling, driver workflows, fleet management, work orders, billing, real-time GPS, and inter-team communications into one app. Before Yardward Pro, a typical haulage business would juggle:")
    add_bullets(doc, [
        "Workforce or similar for time tracking and shift scheduling",
        "Formstack or paper for driver forms (inspections, work orders, dump logs)",
        "Fleetio for vehicle maintenance and history",
        "Twilio plus a spreadsheet for dispatch SMS",
        "QuickBooks Online for invoicing and payroll",
        "Email and group chats for between-role coordination",
    ])
    add_para(doc, "Yardward Pro absorbs every workflow above into a single application with role-specific views — admin, driver, and mechanic — backed by a single Supabase database. The external systems that remain (Twilio, Geotab, QuickBooks, Fleetio, Resend) are integrated, not replaced. Drivers and mechanics in the field still receive Twilio SMS notifications, but the dispatcher composes them from inside Yardward Pro and every reply is tracked back into the right conversation thread.")

    add_heading(doc, "1.1. Architecture at a glance", level=2, color=NAVY)
    add_para(doc, "Yardward Pro is delivered as a Progressive Web App (PWA). That means:")
    add_bullets(doc, [
        "The app loads from a URL in any modern browser — no app-store download or approval cycle",
        "Drivers and mechanics 'install' it from the URL onto their phone home screen with one tap",
        "After install it launches in full-screen, looks and feels native, and updates itself automatically when admins push new releases",
        "Forms submitted in dead zones (rural quarries, basement loading bays, locker rooms) queue locally on the device and sync to the database when the device reconnects",
    ])
    add_para(doc, "The backend is Supabase (managed PostgreSQL with Row-Level Security, Auth, Realtime, Storage, and Edge Functions). Every data access path enforces role-based isolation at the database level — even if a frontend bug tried to leak driver A's time entries to driver B, Postgres itself would refuse to return the rows.")

    add_heading(doc, "1.2. Three role-specific views", level=2, color=NAVY)
    add_para(doc, "Yardward Pro presents a different shell depending on which role the signed-in user has:")

    add_field_table(doc, [
        ("Admin", "Desktop-first sidebar with 17 management tabs covering schedule, fleet, work orders, communications, billing, settings, etc. Used by dispatchers, owners, and office staff."),
        ("Driver", "Mobile-first bottom-tab nav with 6 tabs (Home, Jobs, Forms, Tickets, Messages, Profile). Used in the cab and on-site by truck drivers and equipment operators."),
        ("Mechanic", "Desktop or tablet sidebar with 6 workshop-focused tabs (Dashboard, Work orders, Messages, POs, Maintenance logs, Inventory). Used in the shop by mechanics."),
    ])

    add_callout(doc, 'info', "Role assignment is server-controlled. An admin assigns the role when they create a user. A driver can never elevate themselves to admin — even attempts to manipulate the JWT or localStorage are blocked by RLS at the database level.")

    # ========================================================================
    # SECTION 2 — GETTING STARTED
    # ========================================================================
    add_section_break(doc)
    add_heading(doc, "2. Getting started", level=1, color=AMBER)

    add_heading(doc, "2.1. Sign-in walkthrough", level=2, color=NAVY)
    add_para(doc, "Open https://yardward-pro.vercel.app/ (or your custom domain) in Chrome, Safari, Firefox, or Edge. The sign-in page renders the same on desktop and mobile.")
    add_image(doc, "public-01-login.png", "Yardward Pro sign-in page. Left: hero panel with brand statistics. Right: sign-in form with role chips, email + password, Forgot? link, and Sign in button.")

    add_steps(doc, [
        "Pick your role from the three chips at the top of the form: Admin, Driver, or Mechanic. This is purely a hint — your actual permissions come from the database.",
        "Enter your work email in the Email field. The form accepts any email format that has the shape user@domain.tld; no real-time validation against the database happens at this step.",
        "Enter your password in the Password field. Passwords are masked by default; there is no 'show password' toggle in production to prevent shoulder-surfing.",
        "Click Sign in to Yardward Pro. The button shows a spinner while authentication runs against Supabase Auth.",
        "On success you are redirected to the role-appropriate landing page: /admin for admins, /driver for drivers, /mechanic for mechanics.",
    ])

    add_callout(doc, 'warning', "If you enter the wrong password three times in a row, Supabase Auth applies rate limiting. Wait 30 seconds before retrying. After ten consecutive failures, the account is temporarily locked — contact your admin to unlock via the Supabase dashboard.")

    add_heading(doc, "2.2. First-time sign-in", level=2, color=NAVY)
    add_para(doc, "When an admin creates your account, they hand you a one-time temporary password (a 16-character random string like 'rKszZ@A4!2uCNYJ5'). You must rotate this on first sign-in for security — anyone who saw the temp password during handoff could theoretically log in as you.")

    add_steps(doc, [
        "Sign in using the temp password your admin gave you. You land on your role's dashboard.",
        "Immediately go back to /login (sign out from your profile or just navigate).",
        "Click the Forgot? link next to the Password label.",
        "An inline panel expands with an email input. The form pre-fills with the email you just typed; verify it's yours.",
        "Click Send link. A confirmation toast appears: 'If <email> has an account, a reset link is on its way.'",
        "Open your email and look for a message from noreply@mail.app.supabase.io (or whatever sender you've configured in SMTP).",
        "Click the link in the email. You land on the /reset-password page.",
        "Enter a strong new password (minimum 6 characters — but we strongly recommend 12+ with a mix of uppercase, lowercase, digits, and symbols).",
        "Confirm the password in the second field.",
        "Click Update password. The page redirects to /login after a brief success animation.",
        "Sign in with your new password.",
    ])

    add_image(doc, "public-02-forgot-password.png", "Forgot? panel expanded. Pre-filled with the email already in the sign-in form. Click Send link to fire the reset email.")

    add_heading(doc, "2.3. Forgot password recovery (existing users)", level=2, color=NAVY)
    add_para(doc, "If you've forgotten your password, the recovery flow is identical to first-time sign-in:")
    add_steps(doc, [
        "Go to /login.",
        "Click Forgot?, enter your email, and click Send link.",
        "Check your inbox (and spam folder — until SMTP is configured, the email comes from Supabase's default mailer and is often filtered).",
        "Click the link, set a new password, sign in.",
    ])
    add_callout(doc, 'tip', "Supabase's default mailer is rate-limited to ~4 reset emails per hour across the entire organization. If you don't see the email within 60 seconds, ask your admin to use the Admin → Settings → Users tab to rotate your password directly (no email required).")

    add_heading(doc, "2.4. Installing the mobile PWA", level=2, color=NAVY)
    add_para(doc, "Drivers and mechanics should install Yardward Pro to their phone home screen for the best experience. After install:")
    add_bullets(doc, [
        "The app launches in full-screen, without the browser address bar — feels like a native app",
        "App icon shows your Yardward branding on the home screen",
        "Notifications integrate with the device's push notification system",
        "Forms persist across launches even if you force-close the browser",
    ])

    add_heading(doc, "iOS Safari (iPhone, iPad)", level=3, color=NAVY)
    add_steps(doc, [
        "Open the app URL in Safari (must be Safari — Chrome on iOS cannot install PWAs).",
        "Tap the Share button at the bottom of the screen (square with up-arrow).",
        "Scroll down in the share sheet and tap 'Add to Home Screen'.",
        "Give the shortcut a name (default: 'Yardward Pro').",
        "Tap Add in the top-right.",
        "Close Safari and tap the new icon on your home screen to launch.",
    ])

    add_heading(doc, "Android Chrome", level=3, color=NAVY)
    add_steps(doc, [
        "Open the app URL in Chrome.",
        "Tap the three-dot menu in the top-right.",
        "Tap 'Install app' or 'Add to Home screen' (label depends on Chrome version).",
        "Confirm the install in the dialog.",
        "The app icon appears in your app drawer and on your home screen.",
    ])

    add_callout(doc, 'info', "After install, you may be prompted to allow notifications and location access. Allow both — Yardward Pro uses location to capture GPS on form submissions (required for compliance) and notifications for job assignments and message alerts.")

    add_heading(doc, "2.5. Roles & access overview", level=2, color=NAVY)
    add_para(doc, "Each role has a strict permission boundary enforced by Postgres Row-Level Security. The table below summarizes what each role can see and do:")

    add_field_table(doc, [
        ("Admin", "Full read/write across all tables. Can create/edit/delete users, vehicles, clients, jobs, work orders. Can observe all Communications threads. Can approve work orders and POs. Can configure org settings and integrations. There is no higher role — admins are the ultimate authority in the org."),
        ("Driver", "Read/write only on rows where driver_id = their UUID. Can see only jobs assigned to them, vehicle they're assigned, their own time entries, inspections, work orders, and tickets. Can see clients only through the jobs they're working. Cannot see other drivers' data. Can create Communications threads with mechanics; admins can be tagged."),
        ("Mechanic", "Read/write on shared shop resources (inventory, tools, maintenance logs, vehicle inspections) plus their own purchase requests and assigned maintenance work orders. Cannot see driver-private data (time entries, individual driver schedules). Can create Communications threads with drivers."),
    ])

    # ========================================================================
    # SECTION 3 — ADMIN GUIDE
    # ========================================================================
    add_section_break(doc)
    add_heading(doc, "3. Admin guide", level=1, color=AMBER)
    add_para(doc, "The admin shell is a left-sidebar layout with 17 navigation tabs. The sidebar is fixed on desktop (width 240px) and collapses behind a hamburger menu on tablets and below. The top bar shows the current page title and a notifications bell that pulses when there's unread activity.")

    # 3.1 Dashboard
    add_heading(doc, "3.1. Dashboard", level=2, color=NAVY)
    add_image(doc, "admin-01-dashboard.png", "Admin dashboard — landing page after sign-in.")
    add_para(doc, "The dashboard is your morning check. It summarizes the operational state of the business at a glance and links into every detail screen below.")

    add_heading(doc, "Layout", level=3, color=NAVY)
    add_para(doc, "Three rows of content:")
    add_bullets(doc, [
        "KPI strip — quick numeric summary cards: active drivers, scheduled jobs today, work orders awaiting approval, POs pending, overdue maintenance, etc.",
        "Today's schedule — a horizontally-scrollable view of every job scheduled for today, sorted by start time. Click any card to open the job detail sheet.",
        "Recent activity feed — last 10 notifications from the notifications table, sorted newest-first. Color dots indicate type: green=positive, amber=pending action, red=alert/flag, blue=informational.",
    ])

    add_heading(doc, "Common workflows from the dashboard", level=3, color=NAVY)
    add_bullets(doc, [
        "Click a KPI card to drill into the relevant tab (e.g. 'Work orders awaiting approval' takes you to /admin/work-orders filtered to pending)",
        "Click a job card in the schedule strip to review or edit its details",
        "Click an activity entry to navigate to the source (notification with a link field)",
    ])

    # 3.2 Schedule
    add_heading(doc, "3.2. Schedule", level=2, color=NAVY)
    add_image(doc, "admin-02-schedule.png", "Weekly schedule grid. Rows are drivers; columns are weekdays. Cells show assigned job cards or a + quick-add button.")
    add_para(doc, "The schedule is the dispatcher's primary workspace. It's a weekly grid showing every driver on the y-axis and the 7 days of the current week on the x-axis. Each {driver, day} cell either contains a job card (the driver is assigned a job that day) or a faint + button revealed on hover (the cell is empty and you can quick-add).")

    add_heading(doc, "Creating a new job", level=3, color=NAVY)
    add_para(doc, "Two paths:")
    add_bullets(doc, [
        "Click the 'Create new job' button at the top of the page — opens an empty dialog",
        "Click the + in a specific {driver, weekday} cell — opens the dialog with that driver and date pre-filled",
    ])
    add_image(doc, "admin-02b-create-job-dialog.png", "Create Job dialog. Fields for client, address, date, time, driver, vehicle, and notes.")

    add_heading(doc, "Create Job dialog fields", level=3, color=NAVY)
    add_field_table(doc, [
        ("Client", "Required. Select from your client roster. The client's rate table determines what the invoice will look like once the work order is approved."),
        ("Address", "Free-text job site address. Used for SMS dispatch and the live map. Add a postal code for best results."),
        ("Date", "Job date in YYYY-MM-DD format. Defaults to today (or the weekday you clicked the + cell for)."),
        ("Time", "Scheduled start time in HH:MM 24-hour format. Used for sorting in the schedule grid and the driver's Home screen."),
        ("Driver", "Required. Pick from active drivers. Pre-filled if you opened the dialog from a specific row."),
        ("Vehicle", "Required. Pre-fills with the driver's currently assigned vehicle, but you can override (e.g. for swap-truck days)."),
        ("Notes", "Optional free-text. Internal dispatcher notes — appended to the SMS sent to the driver when you publish."),
    ])

    add_heading(doc, "Save as Draft vs. Publish & notify driver", level=3, color=NAVY)
    add_para(doc, "The dialog has two submit actions:")
    add_bullets(doc, [
        "Save as draft — the job is created with status='draft'. Drafts are private to admins; the driver does not see them and no SMS is sent. Use this when finalizing details over several edits.",
        "Publish + notify driver — the job goes live (status='published'), shows up in the driver's My Jobs and Today's schedule, and a Twilio SMS is dispatched immediately with the job summary, time, address, and any notes.",
    ])

    add_callout(doc, 'warning', "Once a job is published, you can edit it but the driver does NOT get a second SMS. Material changes (different time, different address) should be communicated separately via the Communications tab or by phone.")

    add_heading(doc, "Editing or rescheduling", level=3, color=NAVY)
    add_steps(doc, [
        "Click any existing job card in the grid.",
        "The Edit Job sheet opens on the right with the same fields as Create Job, plus a Status dropdown (Draft / Published / In progress / Completed / Delayed / Cancelled).",
        "Make changes; click Save changes.",
        "If you change the driver or date, the card moves in the grid automatically (realtime via Supabase Realtime subscriptions).",
    ])

    # 3.3 Jobs
    add_heading(doc, "3.3. Jobs", level=2, color=NAVY)
    add_image(doc, "admin-03-jobs.png", "All jobs across all dates, with status filter pills and search.")
    add_para(doc, "The Jobs tab is the flat list view of every job — past, present, and future. Use it when you need to find a specific job or look back at history.")

    add_heading(doc, "Filtering", level=3, color=NAVY)
    add_bullets(doc, [
        "Search box — matches client name, address, driver name, and job ID (substring, case-insensitive)",
        "Status pills — All / Pending / In Progress / Completed. Click to filter; click again to clear",
        "Date range picker (when expanded) — narrow to a billing period or week",
    ])

    add_heading(doc, "Per-row actions", level=3, color=NAVY)
    add_para(doc, "Click any row to open the side-sheet with the job's full lifecycle:")
    add_bullets(doc, [
        "Job details (client, address, scheduled time, vehicle, driver, notes)",
        "Status transitions log (when the job moved Draft → Published → In progress → Completed)",
        "Linked work order if the driver has submitted one",
        "Linked invoice if the work order has been approved",
        "GPS map showing the job site and any clock-in/clock-out points from the driver",
        "Communications threads tagged to this job (filtered by topic_ref_id=job.id)",
    ])

    # 3.4 Drivers & mechanics
    add_heading(doc, "3.4. Drivers & mechanics", level=2, color=NAVY)
    add_image(doc, "admin-04-drivers.png", "Drivers + mechanics roster. Each card shows avatar, name, status badge, phone, license info. Pencil icon opens phone-edit sheet.")
    add_para(doc, "This is your staff roster. Both drivers and mechanics live in this tab, displayed in two sections.")

    add_heading(doc, "The yellow phone-placeholder warning", level=3, color=NAVY)
    add_para(doc, "If any driver or mechanic has a phone number that isn't a valid E.164 (e.g. they have a US format like '555-123-4567' or a blank field), a yellow banner at the top of the page counts them. Twilio outbound SMS will NOT deliver to placeholder numbers — job assignments and Communications messages get silently dropped for those people.")
    add_steps(doc, [
        "Click the pencil icon on a card with a placeholder phone.",
        "A sheet opens with the current phone and a new-phone input.",
        "Enter the real number in E.164 format: + then country code then number, no spaces (e.g. +14165550100 for a Toronto landline).",
        "The Save button only enables when the format is valid AND the new value differs from the current.",
        "Click Save phone. The card updates immediately; the next SMS to that person will deliver.",
    ])

    add_heading(doc, "Adding a new driver", level=3, color=NAVY)
    add_image(doc, "admin-04b-add-driver-dialog.png", "Add Driver dialog. Name, email, E.164 phone, license number, license expiry. The Phone field placeholder is the recommended format.")
    add_para(doc, "Click 'Add driver' at the top-right of the roster.")

    add_field_table(doc, [
        ("Name", "Required. Full name; first name shows in greeting cards, full name shows in lookups and invoices."),
        ("Email", "Required, must be a deliverable address. The driver will sign in with this email. Cannot be changed later without admin intervention."),
        ("Phone (E.164)", "Optional but strongly recommended — without it, the driver cannot receive job-assignment SMS or Communications messages on their phone. Format: +<country><number> with no spaces or punctuation."),
        ("License number", "Required for drivers. Will be displayed on the driver's profile card and used for compliance reporting. Free-text — accepts any format your jurisdiction uses."),
        ("License expiry", "Required for drivers. YYYY-MM-DD. Yardward Pro can be configured to warn the driver and admin N days before expiry."),
    ])

    add_para(doc, "On click 'Add driver', Yardward Pro:")
    add_steps(doc, [
        "Validates the inputs (email format, phone E.164 if provided, license expiry is a real date)",
        "Calls the admin-create-user edge function which uses the Supabase Auth Admin API to create the auth user with a random 16-character temporary password",
        "Inserts a row into the profiles table with the email, name, and role='driver' via the handle_new_auth_user trigger",
        "Patches the phone field onto the new profile",
        "Inserts a row into the drivers side table with license_number, license_expiry, and computed initials",
        "Shows you a 'Driver created' panel with the email, the temp password, and a Copy credentials button",
    ])

    add_callout(doc, 'tip', "Click Copy credentials to put 'Email: <email>\\nTemp password: <pw>' on the clipboard. Paste into a Slack DM, SMS, or email to hand off to the new driver. They use the credentials to sign in once, then immediately rotate via Forgot? as described in section 2.2.")

    # 3.5 Vehicles
    add_heading(doc, "3.5. Vehicles", level=2, color=NAVY)
    add_image(doc, "admin-05-vehicles.png", "Fleet roster. Each card shows vehicle ID, year, type (truck/trailer/equipment), assigned driver, odometer, engine hours, last service date, and next service due date.")
    add_para(doc, "Yardward Pro tracks every vehicle in your fleet — trucks, trailers, equipment, anything you want maintenance tracking on. Each vehicle has:")
    add_bullets(doc, [
        "A unique ID (e.g. TRK-07, EQ-02) — used everywhere as the human-readable identifier",
        "Full make/model/year metadata",
        "An assigned primary driver (optional; one driver per vehicle at a time)",
        "Odometer and engine hours, kept in sync with Geotab telemetry",
        "Last service date and next service due date",
        "Status: operational, in maintenance, or out of service",
    ])

    add_heading(doc, "Adding a vehicle", level=3, color=NAVY)
    add_para(doc, "Click 'Add vehicle' at the top. The dialog accepts ID (auto-generated suggestion T-NN), name, year, type, current odometer, engine hours, and Geotab device ID if you have one to wire up GPS tracking.")

    add_heading(doc, "Importing from Fleetio", level=3, color=NAVY)
    add_para(doc, "If you're migrating off Fleetio, click 'Import from Fleetio'. The import edge function reads from your Fleetio account (via the FLEETIO_BEARER_TOKEN secret) and brings in:")
    add_bullets(doc, [
        "Every vehicle with full metadata",
        "Maintenance history (every service event with parts, labor, cost)",
        "Fuel logs",
    ])
    add_para(doc, "Use 'Dry run' mode first to see what would be imported without writing anything. Then disable dry-run and click Run import — typically takes 30-60 seconds for a 50-vehicle fleet.")

    add_heading(doc, "Vehicle detail page", level=3, color=NAVY)
    add_para(doc, "Click any vehicle card to open its detail page. Tabs:")
    add_bullets(doc, [
        "Overview — current status, assigned driver, embedded GPS map showing where the vehicle is RIGHT NOW",
        "Maintenance — chronological log of every service event, including ones logged from the mechanic shop",
        "Fuel — fill-up log with mpg/L per 100km calculations",
        "Tools — assigned tools (from the tool checklist module)",
        "History — timeline of status changes",
    ])

    # 3.6 Live map
    add_heading(doc, "3.6. Live vehicle map", level=2, color=NAVY)
    add_image(doc, "admin-06-live-map.png", "Live map of every vehicle in the fleet, pulled from Geotab. Color-coded markers by status. Sidebar lists vehicles with last-seen times.")
    add_para(doc, "The live map shows every vehicle's current location, updated every 30 seconds. Markers are color-coded:")
    add_bullets(doc, [
        "Green — operational and on schedule",
        "Amber — in maintenance (currently in the shop)",
        "Gray — out of service",
        "Red — alert (e.g. prolonged stop detected, exceeded route boundary)",
    ])

    add_heading(doc, "How the map gets data", level=3, color=NAVY)
    add_para(doc, "Yardward Pro polls Geotab's API every 5 minutes via a Supabase cron job (the geotab-sync-locations edge function). The poll fetches each vehicle's last known coordinates, last heading, last odometer reading, and engine hours, and writes them into the vehicle_locations table. The map subscribes to that table via Supabase Realtime — so changes show up on your screen within a second of Geotab returning new data.")

    add_heading(doc, "Common workflows", level=3, color=NAVY)
    add_bullets(doc, [
        "Click a marker — opens a popup with vehicle ID, driver name, last-seen time, and 'Open detail' deep-link",
        "Click a sidebar row — recenters the map and zooms to that vehicle",
        "Use the search box (top of sidebar) — find by vehicle ID or driver name",
        "'Refresh now' button (top-right) — forces an immediate Geotab poll instead of waiting for the cron",
    ])

    # 3.7 Work orders
    add_heading(doc, "3.7. Work orders", level=2, color=NAVY)
    add_image(doc, "admin-07-work-orders.png", "Submitted work orders awaiting management review. Columns: WO ID, job, driver, client, weight tonnes, dump site, status, submitted time.")
    add_para(doc, "A work order is the driver's record of work performed at a job site — load type, weight, dump location, foreman signature, and any photos. When a driver completes a job, they submit a work order from their mobile app; it lands here for management review.")

    add_heading(doc, "The approval flow", level=3, color=NAVY)
    add_steps(doc, [
        "Driver submits the work order from /driver/work-order. The submission includes GPS check-in (where the WO was submitted), foreman signature drawn on touchscreen, weight in tonnes, dump site, and any photos.",
        "The WO arrives here with status='pending'. The admin gets a notification.",
        "Click the WO row to open the review sheet. You can see every field, the captured signature, GPS map, and photos.",
        "Click Approve to authorize invoicing. Yardward Pro auto-generates the QuickBooks invoice draft using the client's rate table — the line items are looked up based on the work order's load type and weight. The invoice draft lands in the Forms & Submissions → Invoices tab where you can review before pushing to QBO.",
        "OR click Reject. A dialog asks for a rejection reason (free-text, required). The driver gets a notification with the reason and can resubmit a corrected work order.",
    ])

    add_callout(doc, 'info', "Only one approval action is allowed per work order — the approve_work_order Supabase RPC uses a row-level FOR UPDATE lock to prevent two admins from approving the same WO twice and creating duplicate invoices.")

    # 3.8 Communications
    add_heading(doc, "3.8. Communications", level=2, color=NAVY)
    add_image(doc, "admin-08-communications.png", "Communications inbox. Left rail: thread list with filter chips (Tagged me / Joined / All). Right pane: selected thread with messages and compose box.")
    add_para(doc, "Communications is Yardward Pro's central messaging surface. It implements a customer-specific topology: driver↔mechanic threads are the primary channel; admin has read-everything visibility for oversight, but doesn't get notifications by default. Admins become active participants only when explicitly tagged by a driver or mechanic, or when they self-join a thread to intervene.")

    add_heading(doc, "Filter chips", level=3, color=NAVY)
    add_field_table(doc, [
        ("Tagged me", "DEFAULT. Shows only threads where the admin is an active participant. This is the admin's 'inbox' — threads requiring their attention."),
        ("Joined", "Threads where the admin has self-joined to observe or intervene, but wasn't formally tagged."),
        ("All", "Every conversation in the org — RLS allows admin to read all threads even when not participating. Use to monitor general activity."),
    ])

    add_heading(doc, "Starting a new conversation", level=3, color=NAVY)
    add_image(doc, "admin-08b-new-conversation-dialog.png", "New Conversation dialog (admin path). Topic, subject, participant checkboxes for both drivers and mechanics.")
    add_steps(doc, [
        "Click 'New' at the top-right of the thread list.",
        "Pick a topic from the dropdown: general, job, vehicle, or maintenance.",
        "Enter a subject (one-line summary of what the thread is about).",
        "Check participants from the drivers + mechanics list. You can include multiple — useful for cross-role coordination.",
        "Click 'Open conversation'. The thread is created, you're added as originator, and the right pane switches to the new thread.",
    ])

    add_heading(doc, "Inside a thread", level=3, color=NAVY)
    add_para(doc, "The thread view shows:")
    add_bullets(doc, [
        "Participant pills at the top — initials + role badge for each active member; admins joined-not-tagged show '(observing)' suffix until they post",
        "Message list — scrolls oldest-to-newest, with your messages right-aligned and others' left-aligned",
        "Compose box at the bottom with paperclip (attach photo or PDF), @ button (tag admin — drivers/mechanics only), and send button",
        "Close button (admin or originator only) — archives the thread with optional resolution notes",
    ])

    add_callout(doc, 'tip', "If you're observing a thread (filter = All) and need to intervene, you'll see a 'Join conversation' button instead of the compose box. Click Join, then you become a participant and can post + receive notifications.")

    # 3.9 Timesheets
    add_heading(doc, "3.9. Timesheets", level=2, color=NAVY)
    add_image(doc, "admin-09-timesheets.png", "Weekly timesheet view with clock-in/clock-out events cross-referenced against Geotab vehicle movement.")
    add_para(doc, "Timesheets show every driver's clock-in/clock-out events for the selected week, with automatic cross-reference against Geotab telemetry.")

    add_heading(doc, "How GPS cross-reference works", level=3, color=NAVY)
    add_para(doc, "When a driver clocks in, Yardward Pro records their GPS location (captured by the device). It then asks Geotab where the assigned vehicle was at that exact moment. If the two coordinates are within the GPS tolerance (configurable in Settings → System, default 15 minutes / 200 meters), the entry is OK. If not, the row is flagged.")

    add_heading(doc, "Flag reasons", level=3, color=NAVY)
    add_bullets(doc, [
        "GPS mismatch — driver clocked in from a different location than the truck was",
        "Truck never moved — driver clocked in but the vehicle stayed stationary for the entire shift (possible time-card fraud)",
        "Prolonged stop — the vehicle stopped for >2 hours mid-shift without a corresponding job site visit",
        "Overtime threshold — driver exceeded the configured warning or alert hour threshold",
    ])

    add_heading(doc, "Export to QuickBooks", level=3, color=NAVY)
    add_para(doc, "Click 'Export to QuickBooks' at the top-right. The qbo-push-time edge function batches every approved time entry for the week into QBO TimeActivity records using the configured QBO_REALM_ID. Each driver must have a corresponding QBO Employee mapping configured (Settings → Integrations → QBO mappings).")
    add_callout(doc, 'warning', "Once you've pushed time entries to QBO, you cannot edit them in Yardward Pro — the qbo_payroll_pushes table records which entries were pushed and prevents re-export. If you need to correct a pushed entry, void it in QBO directly.")

    # 3.10 SMS log
    add_heading(doc, "3.10. SMS log", level=2, color=NAVY)
    add_image(doc, "admin-10-sms-log.png", "Every Twilio SMS dispatched with delivery status.")
    add_para(doc, "The SMS log is the audit trail for every Twilio SMS Yardward Pro has sent.")

    add_field_table(doc, [
        ("Sent at", "ISO timestamp of when the SMS was queued with Twilio. The actual delivery happens within seconds but can be delayed by carrier."),
        ("Driver", "Recipient — looked up from drivers.id"),
        ("Job", "Linked job (if the SMS was a job assignment notification)."),
        ("Message", "Body of the SMS. Truncated in the table; click row to see full."),
        ("Twilio ID", "Twilio's message SID (SM...). Useful when contacting Twilio support."),
        ("Delivery", "Current delivery status: queued, sent, delivered, failed. Updated in real-time via Twilio delivery webhooks (when configured)."),
    ])

    add_callout(doc, 'tip', "If a driver claims they never received a job assignment, look here first. A 'failed' status with error code 30003 means the driver's phone is unreachable (powered off, no service). 30007 means the message was blocked as spam by the carrier — common with stock Twilio phone numbers; consider a registered short code for high-volume sends.")

    # 3.11 Purchase orders
    add_heading(doc, "3.11. Purchase orders", level=2, color=NAVY)
    add_image(doc, "admin-11-purchase-orders.png", "PO requests from mechanics awaiting management approval.")
    add_para(doc, "When a mechanic needs a part for a maintenance work order, they submit a PO request from /mechanic. The PO lands here.")

    add_heading(doc, "The inventory check", level=3, color=NAVY)
    add_para(doc, "When a mechanic types a part name in the PO form, Yardward Pro queries the parts inventory in real-time and shows matching items inline. If the part is already in stock with sufficient quantity, the mechanic is encouraged to pull from inventory rather than submit a PO. This prevents over-ordering.")
    add_para(doc, "If the part isn't in stock or the mechanic submits anyway, the request lands here with an inventory snapshot — admin can see whether the suggestion-to-pull-from-stock was offered to the mechanic and overridden.")

    add_heading(doc, "PO statuses", level=3, color=NAVY)
    add_bullets(doc, [
        "pending — submitted by mechanic, awaiting admin review",
        "approved — admin authorized the purchase; mechanic can order",
        "ordered — admin marked as ordered (typically after they've placed the order with the supplier)",
        "received — admin marked as received (parts arrived); the qty_on_hand is bumped via the inventory adjustment",
        "rejected — admin denied with a reason",
    ])

    add_heading(doc, "Approving a PO", level=3, color=NAVY)
    add_para(doc, "Click the PO row to review. The sheet shows the part name, requested qty, cost estimate, the mechanic who submitted, and the maintenance work order it's tied to (if any).")
    add_steps(doc, [
        "If you want to approve, click Approve. The PO status becomes 'approved' and the mechanic is notified.",
        "If you've already placed the order with the supplier, click 'Mark ordered'. The PO becomes 'ordered'.",
        "When parts arrive, open the PO again and click 'Mark received'. Yardward Pro adjusts the inventory qty_on_hand automatically.",
        "OR click Reject. A reason dialog appears. The mechanic gets a notification with the reason; they can resubmit if appropriate.",
    ])

    # 3.12 Prepaid tickets
    add_heading(doc, "3.12. Prepaid tickets", level=2, color=NAVY)
    add_image(doc, "admin-12-prepaid-tickets.png", "Per-client prepaid ticket balances.")
    add_para(doc, "Prepaid tickets is a billing model used in waste hauling and similar industries: a client buys tickets in bulk (e.g. 100 tickets at a flat per-ticket rate), and the driver pulls one ticket per load. The Prepaid Tickets tab shows every client's current balance.")

    add_field_table(doc, [
        ("Client", "Linked client."),
        ("Balance", "Current ticket count. Decrements when drivers record ticket pulls in /driver/tickets."),
        ("Threshold", "Low-balance trigger. When the balance hits this value, Yardward Pro notifies admin to start a replenishment."),
        ("Bundle size", "How many tickets are in one purchase bundle (typical: 50 or 100)."),
        ("Bundle price", "Price per bundle. Used to auto-generate the replenishment invoice."),
        ("Auto-billing", "Boolean. If on, Yardward Pro generates the replenishment invoice automatically when the threshold is hit. If off, admin manually triggers."),
        ("Report frequency", "How often the client receives a usage report email: off / daily / weekly / monthly."),
    ])

    # 3.13 Clients
    add_heading(doc, "3.13. Clients", level=2, color=NAVY)
    add_image(doc, "admin-13-clients.png", "Client roster with contact info and rate tables.")
    add_para(doc, "The clients tab is your customer roster. Each client has a contact (name + email + phone), a billing address, an optional rate table that drives invoice line items, and optional prepaid ticket settings.")

    add_heading(doc, "Creating a new client", level=3, color=NAVY)
    add_image(doc, "admin-13b-new-client-dialog.png", "New Client dialog. Company name, primary contact, email, phone, billing address, notes.")
    add_steps(doc, [
        "Click 'New client' at the top-right.",
        "Fill in Company name (required), Primary contact name, Email (optional but recommended for invoice delivery), Phone, Billing address, Notes.",
        "Click Create client. The client is added to the roster; you can now schedule jobs against them.",
    ])

    add_heading(doc, "Rate tables", level=3, color=NAVY)
    add_para(doc, "Click a client row to open their detail sheet. The Rate table section lets you define custom pricing per load type or vehicle type:")
    add_bullets(doc, [
        "Each line item specifies: load type (e.g. 'Topsoil', 'Recycled concrete', 'General waste'), vehicle type (truck/trailer/equipment), unit (per tonne / per load / per hour), and unit price",
        "When a work order is approved against this client, Yardward Pro looks up the matching line item and generates the QBO invoice with the right qty * unit price",
        "Falls back through the unit chain: per-tonne match first, then per-load, then per-hour",
    ])

    # 3.14 Forms & submissions
    add_heading(doc, "3.14. Forms & submissions", level=2, color=NAVY)
    add_image(doc, "admin-14-forms.png", "Centralized inbox for every form a driver submits.")
    add_para(doc, "Forms & Submissions is the central inbox for everything drivers fill out on the mobile app. The page has tab pills filtering by form type:")
    add_bullets(doc, [
        "Inspections — vehicle pre-trip and end-of-trip safety inspections",
        "Tool checklists — start and end of shift tool inventory",
        "Work orders — on-site work-order submissions awaiting management review",
        "Start of day — clock-in events with GPS",
        "End of day — clock-out events",
        "Job logs — free-form notes during a job",
        "Ticket photos — driver-uploaded photos of paper tickets pulled for prepaid clients",
    ])
    add_para(doc, "Each submission row shows: driver, vehicle (where applicable), submitted timestamp, GPS coordinates if captured (clickable, opens a map), and a Resolve button if the submission is informational. Click any row to open the full detail sheet with all answers, photos, signature (if collected), and the GPS map.")

    # 3.15 Tenders
    add_heading(doc, "3.15. Tenders", level=2, color=NAVY)
    add_image(doc, "admin-15-tenders.png", "Aggregated tender feed scraped from municipal portals.")
    add_para(doc, "The Tenders tab is an aggregated feed of contract opportunities scraped from configured municipal portals. By default Yardward Pro is configured for Halton region (Ontario), but you can add additional sources via the tender_sources table.")

    add_heading(doc, "How the scraper runs", level=3, color=NAVY)
    add_para(doc, "A Supabase cron job runs the tender-scrape edge function every Monday at 06:00 UTC. The function loops over enabled tender_sources, scrapes each portal's tender listings, and inserts new rows into the tenders table.")

    add_heading(doc, "Weekly digest email", level=3, color=NAVY)
    add_para(doc, "After scraping, the function compiles a digest email summarizing new tenders from the past week and sends it to TENDER_DIGEST_RECIPIENTS via Resend. The email is plain text with title, source, deadline, and link to the original tender. Configure recipients in Settings → Integrations → Tender digest.")
    add_callout(doc, 'info', "If RESEND_API_KEY is not configured, the scraper still runs but skips the email send with a warning logged to error_log.")

    add_heading(doc, "On-demand controls", level=3, color=NAVY)
    add_bullets(doc, [
        "'Run scraper now' — fires the scrape function immediately. Useful when you've just added a new source or want to refresh outside the weekly cadence.",
        "'Send test digest' — composes and sends the digest email to the configured recipients on demand. Useful for verifying SMTP setup.",
    ])

    # 3.16 Error log
    add_heading(doc, "3.16. Error log", level=2, color=NAVY)
    add_image(doc, "admin-16-errors.png", "Error log with severity, source, code, message, and resolve action. Includes Dead-letter queue tab.")
    add_para(doc, "The error log captures every notable failure across the app:")
    add_bullets(doc, [
        "React error boundaries — uncaught exceptions in the frontend",
        "Edge Function failures — webhook signature mismatches, Twilio errors, QBO push errors",
        "RPC errors — Supabase SECURITY DEFINER function exceptions",
        "Webhook handler errors — anything that fell into the catch-all outer try",
        "Inbound SMS that couldn't be routed (no matching profile by phone)",
    ])
    add_para(doc, "Each row has a severity (info / warning / error / critical), source, error code (machine-readable), message (human-readable), context (JSON blob with relevant IDs), stack trace (for code errors), and the user who triggered it.")

    add_heading(doc, "Resolving errors", level=3, color=NAVY)
    add_para(doc, "Click any row to open the detail sheet. The Resolve action marks the error as resolved with optional notes. Resolved errors filter out of the main view (toggleable). Use the Severity filter to focus on critical/error rows first.")

    add_heading(doc, "Dead-letter queue", level=3, color=NAVY)
    add_para(doc, "The Dead-letter queue tab shows form submissions that exhausted their retry budget. When a driver submits a form offline and the device retries the sync over and over without success, eventually the submission lands here. Click 'Requeue' to attempt one more sync; if the underlying issue (e.g. RLS violation, malformed data) isn't fixed it will fail again.")

    # 3.17 Reports
    add_heading(doc, "3.17. Reports", level=2, color=NAVY)
    add_image(doc, "admin-17-reports.png", "Operational reports menu — cards for each report type.")
    add_para(doc, "Yardward Pro includes operational reports for management decision-making:")
    add_bullets(doc, [
        "Hours summary — weekly/monthly driver hours, with flagged events broken out",
        "Vehicle utilization — % time operational vs. in maintenance vs. out of service",
        "Revenue by client — invoice totals per client over selected period",
        "Flagged events — every flag of any kind (GPS mismatches, tool damage, overtime warnings)",
        "Maintenance costs — total $ spent on parts and labor per vehicle",
        "Tender pipeline — submitted bids and win/loss outcomes",
    ])
    add_para(doc, "Click any report card to open. Each report has a date-range picker, optional filters (client, driver, vehicle), and a chart + data table. Click 'Export PDF' or 'Export CSV' to download for sharing or further analysis.")

    # 3.18 Settings
    add_heading(doc, "3.18. Settings", level=2, color=NAVY)
    add_para(doc, "The Settings tab has six sub-tabs covering all org-level configuration.")

    add_heading(doc, "Organization profile", level=3, color=NAVY)
    add_image(doc, "admin-18-settings-org.png", "Organization tab — business name, tax ID, address, timezone, currency.")
    add_para(doc, "These fields render across the app: header brand, invoice header, timezone-aware date displays, currency symbol on financial reports. Set them once during onboarding and update only when business details change. Click 'Save changes' — the values persist to the app_settings singleton row in Supabase.")

    add_heading(doc, "System thresholds", level=3, color=NAVY)
    add_image(doc, "admin-18b-settings-system.png", "System tab — GPS tolerance, overtime warning/alert hours, inspection min/max duration.")
    add_field_table(doc, [
        ("GPS tolerance (minutes)", "How wide a window Yardward Pro accepts between driver's clock-in GPS and the vehicle's GPS at the same moment. Default 15. Lower = stricter."),
        ("Overtime warning (hours)", "Driver hours threshold at which a warning notification fires. Default 40."),
        ("Overtime alert (hours)", "Driver hours threshold at which a critical alert fires. Default 44. Must be greater than warning."),
        ("Inspection min duration (seconds)", "Minimum time a pre-trip inspection should take. Faster than this = likely rushed. Default 780 (13 min)."),
        ("Inspection max duration (seconds)", "Maximum reasonable inspection time. Slower = likely distracted. Default 1200 (20 min)."),
    ])

    add_heading(doc, "Integrations", level=3, color=NAVY)
    add_image(doc, "admin-18c-settings-integrations.png", "Integrations tab — Twilio, Geotab, QBO, Fleetio, Resend connection cards.")
    add_para(doc, "Each integration has a connection card showing current status (connected / not connected / error) and Test / Connect / Disconnect buttons. Configuration values come from Supabase function secrets (set via supabase secrets set), not from this UI — these cards only show status and let you trigger health checks.")

    add_heading(doc, "Users", level=3, color=NAVY)
    add_image(doc, "admin-18d-settings-users.png", "Users tab — roster of all admins, mechanics, drivers with Invite user button.")
    add_para(doc, "The Users tab lists every active user in the org and lets admins invite new ones. Click 'Invite user' to open a dialog with email, name, optional phone, role selector (admin/driver/mechanic), and conditional fields. If role=driver, two extra fields appear: license number and license expiry. On submit, Yardward Pro creates the auth user via the admin-create-user edge function and shows you the credentials panel for handoff.")

    add_heading(doc, "Notifications", level=3, color=NAVY)
    add_image(doc, "admin-18e-settings-notifications.png", "Notifications tab — org-wide toggle switches for each notification type.")
    add_para(doc, "Org-wide defaults for notification types: new job assigned SMS, work order awaiting approval, tool flagged on checklist, GPS mismatch on time entry, PO awaiting approval, vehicle maintenance overdue, daily summary email. Drivers and mechanics can override these per-user via their profile.")

    add_heading(doc, "Billing", level=3, color=NAVY)
    add_image(doc, "admin-18f-settings-billing.png", "Billing tab — subscription plan, renewal date, seats used, vehicles active, Cancel subscription button.")
    add_para(doc, "Shows your current subscription state. Plan name, renewal date, seats used vs. limit, active vehicles vs. limit. The 'Cancel subscription' button opens a confirm dialog with an optional reason; on submit it flips billing_status to 'cancel-requested' and notifies the Yardward team. Your subscription remains active until the next renewal date; you'll receive a follow-up to confirm.")

    # ========================================================================
    # SECTION 4 — DRIVER GUIDE
    # ========================================================================
    add_section_break(doc)
    add_heading(doc, "4. Driver guide (mobile)", level=1, color=AMBER)
    add_para(doc, "The driver shell is mobile-first. The bottom navigation bar has 6 tabs and stays visible at all times. Forms are designed for one-handed use with large tap targets — drivers wearing gloves or working in poor light can still hit the right buttons.")

    # 4.1 Home
    add_heading(doc, "4.1. Home", level=2, color=NAVY)
    add_image(doc, "driver-01-home.png", "Driver home screen — today's jobs, clock-in button, GPS status badge, quick links to forms.", width_in=3.5)
    add_para(doc, "The home screen is your morning check. It shows:")
    add_bullets(doc, [
        "GPS status badge at the top — green if active and accurate, amber if in fallback mode (using last known position), gray if pending",
        "Clock-in / Clock-out button — big touch target. Tap to start your shift",
        "Pre-trip inspection lockout indicator — if you haven't done your pre-trip in the last 12 hours, you can't clock in (CVOR rule)",
        "Today's jobs list — each card shows scheduled time, client, address, truck",
        "Quick links — shortcuts to Start of day, Tool checklist, Inspection, Work order",
    ])

    # 4.2 My jobs
    add_heading(doc, "4.2. My jobs", level=2, color=NAVY)
    add_image(doc, "driver-02-jobs.png", "Today's and upcoming jobs assigned to you, sorted by scheduled time.", width_in=3.5)
    add_para(doc, "The Jobs tab shows your full schedule. Today's jobs are highlighted at the top; future jobs are below. Tap any job to open its details:")
    add_bullets(doc, [
        "Client name and full address (tappable — opens in your phone's map app for directions)",
        "Scheduled start time",
        "Assigned vehicle",
        "Dispatcher notes (anything the admin added during scheduling)",
        "Submit work order button — opens /driver/work-order with the job context pre-filled",
    ])

    # 4.3 Forms tile menu
    add_heading(doc, "4.3. Forms tile menu", level=2, color=NAVY)
    add_image(doc, "driver-03-forms.png", "Forms tile menu — large tiles for each driver form.", width_in=3.5)
    add_para(doc, "The Forms tab is a tile menu — large tappable tiles for each form. From here you can jump into any of the driver workflows below.")

    # 4.4 Start of day
    add_heading(doc, "4.4. Start of day & clock-in", level=2, color=NAVY)
    add_image(doc, "driver-04-start-of-day.png", "Start of day form — vehicle selection, GPS captured, optional notes.", width_in=3.5)
    add_para(doc, "Start of day clocks you in for the shift.")
    add_steps(doc, [
        "Tap Forms → Start of day (or tap the big Clock-in button on Home).",
        "Confirm the vehicle you're driving (pre-selected if you have an assigned truck, but overridable).",
        "Wait for the GPS badge to turn green — this captures your current location for cross-reference against Geotab telemetry.",
        "Optional: add a note (e.g. 'Late start due to weather').",
        "Tap Submit. The form is saved; you're clocked in.",
    ])
    add_callout(doc, 'warning', "If you haven't completed a passing pre-trip inspection in the last 12 hours, the Submit button is disabled with a banner: 'Complete inspection first'. Tap the banner to jump to /driver/inspection.")

    # 4.5 Tool checklist
    add_heading(doc, "4.5. Tool checklist", level=2, color=NAVY)
    add_image(doc, "driver-05-tool-checklist.png", "Tool checklist — per-vehicle tool inventory with OK / Flag toggles.", width_in=3.5)
    add_para(doc, "The tool checklist verifies your truck's tool inventory at start and end of shift. Each tool has three possible states:")
    add_bullets(doc, [
        "OK (default) — tool is present and in working condition",
        "Missing — tool is not in the truck",
        "Damaged — tool is present but broken",
    ])
    add_para(doc, "Any item flagged Missing or Damaged requires a note explaining the issue. Yardward Pro notifies admin and the assigned mechanic immediately so they can dispatch a replacement or schedule a repair.")
    add_callout(doc, 'tip', "You can't complete End of day without a passing end-of-shift tool checklist. Get into the habit of running the checklist at the start of every shift (so any missing tool can be addressed before you leave the yard) and again before you submit End of day.")

    # 4.6 Vehicle inspection
    add_heading(doc, "4.6. Vehicle inspection (pre-trip)", level=2, color=NAVY)
    add_image(doc, "driver-06-inspection.png", "Pre-trip safety inspection — tyres, lights, brakes, etc. Photo capture for flagged items.", width_in=3.5)
    add_para(doc, "The pre-trip inspection is required by CVOR (Commercial Vehicle Operator's Registration) before every shift. Yardward Pro guides you through 8 inspection categories with a checklist of items per category.")

    add_heading(doc, "Inspection categories", level=3, color=NAVY)
    add_bullets(doc, [
        "Tyres & wheels — tread depth, sidewall damage, lug nuts",
        "Lights & indicators — headlights, brake lights, signals, hazards",
        "Brakes — feel of pedal, parking brake, air pressure if applicable",
        "Fluid levels — engine oil, coolant, brake fluid, transmission, washer fluid",
        "Signage & decals — DOT numbers, CVOR sticker, hazard signs",
        "Coupling devices (if towing) — kingpin, fifth wheel, safety chains",
        "Mirrors & windows — cracks, visibility, wiper function",
        "Cab equipment — seatbelts, fire extinguisher, first aid kit",
    ])

    add_heading(doc, "Submitting the inspection", level=3, color=NAVY)
    add_steps(doc, [
        "Walk around the truck doing the physical inspection.",
        "In the app, tap each item to mark it OK or flag it with a defect category.",
        "Flagged items require: a note describing the issue AND a photo (auto-attached from the camera).",
        "The form saves drafts every 5 seconds — if you lose signal mid-form, your progress is preserved.",
        "When all items are checked, tap Submit. The form captures: timestamp, GPS, inspection duration, vehicle, all answers, all photos.",
    ])
    add_callout(doc, 'info', "A passing inspection unlocks the 12-hour drive window. The system tracks last_pretrip_at on the vehicle. If you try to clock in or submit a work order more than 12 hours after the last passing pretrip, you're locked out until you redo the inspection.")
    add_callout(doc, 'warning', "If you flag any item as Critical (brakes, tyres, fluid leaks), the inspection is failed and the vehicle is auto-flagged out_of_service. Admin and mechanic are notified immediately. Do NOT drive a failed-inspection vehicle.")

    # 4.7 Job log
    add_heading(doc, "4.7. Job log", level=2, color=NAVY)
    add_image(doc, "driver-07-job-log.png", "Job log — free-form notes during a shift. Anchored to current job.", width_in=3.5)
    add_para(doc, "The job log is your dictation surface. Use it to record anything noteworthy during the shift — delay reasons, client disputes, equipment issues, conversations with customers. Each entry is timestamped + GPS-tagged automatically.")
    add_para(doc, "Use the job log entries later when:")
    add_bullets(doc, [
        "Disputing a customer complaint ('I told them at 10:23 AM that we needed access to the back lot')",
        "Explaining to admin why a job took longer than estimated",
        "Documenting an incident for compliance records",
        "Tracking patterns (e.g. a recurring issue with a specific client)",
    ])

    # 4.8 Work order
    add_heading(doc, "4.8. Work order submission", level=2, color=NAVY)
    add_image(doc, "driver-08-work-order.png", "Work order form — load type, weight, dump site, foreman signature, photos, GPS check-in.", width_in=3.5)
    add_para(doc, "The work order is the most important form in the driver workflow. Submit one per completed delivery. After admin approval, the invoice auto-generates and bills the client.")

    add_heading(doc, "Work order fields", level=3, color=NAVY)
    add_field_table(doc, [
        ("Job", "Required. Select the job this work order is for. Pre-filled if you came from /driver/jobs."),
        ("Load type", "Required. Picker — list of load types is configured per-client in their rate table."),
        ("Weight (tonnes)", "Required. Manual entry. Decimal allowed (e.g. 12.4). In a future release, will integrate with scale-house ticket photos."),
        ("Dump site", "Required. Auto-populated from job location but editable."),
        ("Work performed", "Free-text description of the work."),
        ("Foreman signature", "Required. The on-site foreman signs the touchscreen. Stored as a PNG."),
        ("Photo", "Optional but recommended. A photo of the delivered load or the dump site. Multiple allowed."),
        ("GPS check-in", "Automatic. Captures the GPS coordinates where the form was submitted. Used for cross-reference if the work order is disputed."),
    ])

    add_heading(doc, "After submission", level=3, color=NAVY)
    add_steps(doc, [
        "The work order lands in admin's Work Orders tab with status='pending'.",
        "Admin reviews and clicks Approve or Reject.",
        "If approved, Yardward Pro auto-generates the QuickBooks invoice using the client's rate table (matching load_type and weight to a rate line item).",
        "The invoice draft sits in the Invoices tab; admin reviews and pushes to QBO when ready.",
        "If rejected, you get a notification with the reason and can resubmit a corrected work order.",
    ])

    # 4.9 End of day
    add_heading(doc, "4.9. End of day & clock-out", level=2, color=NAVY)
    add_image(doc, "driver-09-end-of-day.png", "End of day form — end-of-shift tool checklist confirmation, fuel level, odometer, vehicle return location.", width_in=3.5)
    add_para(doc, "End of day clocks you out for the day.")
    add_steps(doc, [
        "Complete an end-of-shift tool checklist first (Forms → Tool checklist).",
        "Tap End of day.",
        "Confirm the vehicle return location (yard, depot, jobsite — pre-filled).",
        "Enter the final odometer reading.",
        "Optional: fuel level (if you don't refuel and want to flag low fuel for the next driver).",
        "Optional: end-of-shift notes.",
        "Tap Submit. The form is saved; you're clocked out.",
    ])
    add_callout(doc, 'warning', "If you skip the tool checklist, the End of day form blocks submission. Yardward Pro tracks tool_checklist_submissions with a kind field ('start_of_shift' vs 'end_of_shift') and refuses to clock you out without the matching end-of-shift submission.")

    # 4.10 Tickets
    add_heading(doc, "4.10. Prepaid tickets", level=2, color=NAVY)
    add_image(doc, "driver-10-tickets.png", "Record a prepaid ticket pull — client picker, qty entry, optional photo.", width_in=3.5)
    add_para(doc, "Use this form when you're working a client on prepaid tickets (e.g. waste hauling). Record one ticket pull per load.")
    add_steps(doc, [
        "Tap Tickets in the bottom nav.",
        "Pick the client from the dropdown.",
        "Enter the quantity of tickets pulled (usually 1).",
        "Optional: tap the camera icon to capture a photo of the physical ticket (proof for billing disputes).",
        "Tap Record. The client's balance auto-decrements; the photo (if attached) is stored in Storage.",
    ])
    add_para(doc, "When the client's balance hits the configured threshold, admin gets a low-balance notification and either auto-bills a replenishment (if auto-billing is on for that client) or manually triggers one.")

    # 4.11 Messages
    add_heading(doc, "4.11. Messages", level=2, color=NAVY)
    add_image(doc, "driver-11-messages.png", "Driver's view of Communications — thread list and conversation Sheet.", width_in=3.5)
    add_para(doc, "Messages is the driver's view of Communications. You can:")
    add_bullets(doc, [
        "See every conversation you're part of",
        "Start a new thread with a mechanic (pick recipient, set topic, enter subject)",
        "Tag an admin into an existing thread when you need their input (tap the @ button in the compose box)",
        "Attach photos to messages (paperclip button)",
    ])
    add_para(doc, "Messages flow over real Twilio SMS even when you're outside the app. The mechanic gets a text on their phone and their reply comes back into the same thread within seconds. This is the customer-confirmed topology: driver↔mechanic is the primary channel, with admin tagged when oversight is needed.")

    # 4.12 Profile
    add_heading(doc, "4.12. Profile", level=2, color=NAVY)
    add_image(doc, "driver-12-profile.png", "Profile screen — name, license info, shift status, action rows.", width_in=3.5)
    add_para(doc, "The Profile screen shows your account details and four action rows:")
    add_field_table(doc, [
        ("Change password", "Fires a password reset email to your address. Click the link in the email and set a new password."),
        ("Notifications", "Opens a sheet with 6 toggle switches for per-user notification preferences. Override the org-wide defaults set by admin."),
        ("Help & support", "Opens a sheet with FAQ accordion, mailto/tel links to support, and a ticket form. Tickets go to support_tickets table for admin triage."),
        ("Logout", "Signs out of the app. You'll be redirected to /login."),
    ])

    # ========================================================================
    # SECTION 5 — MECHANIC GUIDE
    # ========================================================================
    add_section_break(doc)
    add_heading(doc, "5. Mechanic guide", level=1, color=AMBER)
    add_para(doc, "The mechanic shell is a tablet/desktop sidebar layout focused on the workshop queue. Six tabs cover everything a mechanic does day-to-day.")

    add_heading(doc, "5.1. Workshop dashboard", level=2, color=NAVY)
    add_image(doc, "mechanic-01-dashboard.png", "Mechanic dashboard — active MWOs assigned to me, PO submit form, low-stock alerts.")
    add_para(doc, "The dashboard is your queue. Active maintenance work orders assigned to you are at the top, with vehicle ID, priority, issue summary, and the driver who reported it. Use the inline PO request form to submit a new parts request.")

    add_heading(doc, "5.2. Work orders queue", level=2, color=NAVY)
    add_image(doc, "mechanic-02-work-orders.png", "All maintenance work orders sorted by priority then creation date. Claim button for unclaimed MWOs.")
    add_para(doc, "Maintenance work orders (MWOs) come from three sources:")
    add_bullets(doc, [
        "Auto-generated from flagged vehicle inspections — when a driver flags a critical defect, an MWO is auto-created against the vehicle with priority=critical",
        "Driver-submitted — drivers can manually report maintenance issues via /driver/inspection's Report issue button",
        "Admin-created — admin can add maintenance work directly via /admin/vehicles → Add record",
    ])

    add_heading(doc, "Claiming an MWO", level=3, color=NAVY)
    add_para(doc, "Unclaimed MWOs in the queue have a Claim button. Tap it to assign the MWO to yourself. The claim_maintenance_work_order RPC uses a Postgres FOR UPDATE lock to prevent two mechanics from claiming the same MWO simultaneously — only one wins; the other gets an error.")

    add_heading(doc, "Working an MWO", level=3, color=NAVY)
    add_steps(doc, [
        "After claiming, the MWO moves to your active queue.",
        "Open it to see full details: vehicle, reported issue, photos from the driver, priority, source inspection (if applicable).",
        "Update progress: add parts used, log labor hours, attach photos of the work.",
        "When finished, tap Complete. The vehicle's status moves back to operational, and last_service / next_service_due are updated.",
        "The MWO history is preserved in maintenance_logs forever — useful for warranty claims and CVOR audits.",
    ])

    add_heading(doc, "5.3. Messages", level=2, color=NAVY)
    add_image(doc, "mechanic-03-messages.png", "Mechanic's Communications view. Default recipient when starting a new thread is the driver from the most recent MWO.")
    add_para(doc, "Same Communications surface as drivers, with one tweak: when you start a new conversation, the recipient picker defaults to the driver who reported your most recent claimed MWO. This is the most common case — 'I need to ask the driver who reported this fault for more details.'")

    add_heading(doc, "5.4. Purchase requests", level=2, color=NAVY)
    add_image(doc, "mechanic-04-purchase-requests.png", "Your submitted PO requests with status. Use the inline form on the dashboard or the New PO button here.")
    add_para(doc, "Submit a PO request when you need a part that isn't in inventory. The form has:")
    add_field_table(doc, [
        ("Item needed", "Free-text. As you type, Yardward Pro searches the parts inventory in real-time and shows matches inline. If a match exists with sufficient qty_on_hand, you're prompted to pull from stock instead."),
        ("Quantity", "How many of the item you need."),
        ("Cost estimate", "Your best guess at the unit cost. Helps admin approval decisions."),
        ("Justification", "Brief explanation. The maintenance work order this is tied to, or why the part is needed."),
    ])
    add_para(doc, "On submit, the PO lands in admin's Purchase Orders tab with status='pending'. You'll get a notification when it's approved, rejected, or marked received.")

    add_heading(doc, "5.5. Maintenance logs", level=2, color=NAVY)
    add_image(doc, "mechanic-05-maintenance.png", "Historical maintenance log across all vehicles. Used for compliance and warranty.")
    add_para(doc, "Maintenance logs is the permanent record of every maintenance event across the entire fleet. Each entry shows:")
    add_bullets(doc, [
        "Vehicle, date, mechanic who did the work",
        "Type of maintenance (preventive, corrective, emergency)",
        "Description of work performed",
        "Parts used (with qty and cost)",
        "Labor hours",
        "Total cost",
        "Originating MWO (if any)",
    ])
    add_para(doc, "Click any row to see full details. This data is your audit trail for CVOR compliance and warranty claims — never deleted, only resolved.")

    add_heading(doc, "5.6. Parts inventory", level=2, color=NAVY)
    add_image(doc, "mechanic-06-inventory.png", "Parts inventory — quantity on hand, qty reserved for in-progress MWOs, reorder point.")
    add_para(doc, "The parts inventory shows your stock levels:")
    add_field_table(doc, [
        ("Item name", "Part description (e.g. 'Brake pad set — front, Isuzu FXZ')."),
        ("SKU", "Optional manufacturer or internal SKU."),
        ("Qty on hand", "Physical stock currently in the shop."),
        ("Qty reserved", "Stock allocated to in-progress MWOs but not yet pulled. Increments when an MWO references the part."),
        ("Reorder point", "When qty_on_hand - qty_reserved drops below this, admin gets a low-stock notification."),
        ("Unit cost", "Last known cost. Updated when POs are marked received."),
    ])
    add_para(doc, "Click a row to update qty_on_hand (after receiving a delivery) or edit metadata. Use the search box to find a part by name or SKU.")

    # ========================================================================
    # SECTION 6 — TOKENIZED LINKS
    # ========================================================================
    add_section_break(doc)
    add_heading(doc, "6. Tokenized driver links (QR access)", level=1, color=AMBER)
    add_para(doc, "Tokenized driver links let you give a driver one-shot access to a specific form without them needing an account. Common use cases:")
    add_bullets(doc, [
        "A subcontractor driver hauling a single load for you — no need to create them an account",
        "A delivery driver who just needs to record a single prepaid ticket pull",
        "A one-time scale-house operator who needs to confirm receipt of a load",
        "An emergency replacement driver covering a sick employee's shift",
    ])

    add_heading(doc, "6.1. Generating a tokenized link", level=2, color=NAVY)
    add_steps(doc, [
        "Go to Settings → Driver tokens.",
        "Click 'Generate token'.",
        "Pick the scope (which form the token grants access to): start_of_day, tickets, work_order, inspection, etc.",
        "Optionally tie it to a specific driver name (for the audit log) or job ID (for context).",
        "Set the expiry (default: 24 hours).",
        "Click Generate. Yardward Pro creates a 256-bit random token via the gen_random_bytes function and stores it in the driver_tokens table.",
        "The dialog shows you the shareable URL: /t/<token>.",
        "Click 'Copy URL' to put it on the clipboard. Send via SMS, email, or QR code.",
    ])

    add_heading(doc, "6.2. How the link works for the driver", level=2, color=NAVY)
    add_para(doc, "When the driver opens /t/<token> on their phone:")
    add_steps(doc, [
        "The /t/$token route validates the token via the validate_driver_token RPC.",
        "If valid (not expired, not already burned), it establishes a sessionStorage-backed driver-token session.",
        "The driver is redirected to the scoped path (e.g. /driver/tickets).",
        "They complete and submit the form — no login screen, no password.",
        "On successful submission, the token is auto-burned (status='used'); future visits to the URL get 'Token expired'.",
    ])
    add_callout(doc, 'info', "Driver-token sessions are scope-locked. A token minted for /driver/tickets cannot be used to access /driver/work-order or any other route — the route guards check the scope and bounce mismatched paths.")

    # ========================================================================
    # SECTION 7 — OFFLINE HANDLING
    # ========================================================================
    add_section_break(doc)
    add_heading(doc, "7. Offline handling", level=1, color=AMBER)
    add_para(doc, "Drivers often work in places with poor or no cell signal — rural quarries, basement loading bays, locker rooms, etc. Yardward Pro handles this gracefully.")

    add_heading(doc, "7.1. How offline submission works", level=2, color=NAVY)
    add_steps(doc, [
        "Driver fills out a form (e.g. work order) and taps Submit.",
        "Yardward Pro detects there's no internet connection.",
        "Instead of failing, the submission is queued locally in IndexedDB via the offline-queue module.",
        "A small badge appears in the corner of the screen indicating N pending submissions.",
        "The driver continues working — the form is treated as submitted for all UI purposes (e.g. they can proceed to the next step).",
        "When the device reconnects, the queue auto-flushes — every queued submission is replayed against the API in the order it was created.",
        "On successful flush, the badge clears.",
    ])

    add_heading(doc, "7.2. Idempotency and retry safety", level=2, color=NAVY)
    add_para(doc, "Every queued submission carries a client-generated idempotency key. The server-side INSERT uses a UNIQUE constraint on (sender_id, idempotency_key) — even if the network blips and the device retries the same submission, the database refuses to insert a duplicate. No risk of double-submitted work orders or duplicate messages.")

    add_heading(doc, "7.3. Dead-letter queue", level=2, color=NAVY)
    add_para(doc, "If a submission fails permanently (e.g. the server rejects it due to RLS or malformed data), it lands in the dead-letter queue after exhausting its retry budget. Admin can review and requeue from /admin/errors → Dead-letter tab.")

    # ========================================================================
    # SECTION 8 — NOTIFICATIONS
    # ========================================================================
    add_section_break(doc)
    add_heading(doc, "8. Notifications & alerts", level=1, color=AMBER)
    add_para(doc, "Yardward Pro fires notifications for every event that requires someone's attention.")

    add_heading(doc, "8.1. Notification channels", level=2, color=NAVY)
    add_bullets(doc, [
        "In-app notifications — populate the notifications table; render in the bell badge at the top of every page (click to see list)",
        "Twilio SMS — for events where the user is likely on the road (job assignments, urgent alerts)",
        "Email (via Resend) — for digest-style summaries (daily/weekly reports)",
    ])

    add_heading(doc, "8.2. Notification types", level=2, color=NAVY)
    add_field_table(doc, [
        ("new_job_assigned_sms", "Sent to driver when admin publishes a new job assignment."),
        ("work_order_awaiting_approval", "Sent to admin when driver submits a work order."),
        ("tool_flagged_on_checklist", "Sent to admin + mechanic when driver flags a missing/damaged tool."),
        ("gps_mismatch_on_time_entry", "Sent to admin when timesheet cross-reference detects a discrepancy."),
        ("po_awaiting_approval", "Sent to admin when mechanic submits a PO request."),
        ("vehicle_maintenance_overdue", "Sent to admin + mechanic when a vehicle exceeds its next_service_due date."),
        ("daily_summary_email", "Daily digest email summarizing the day's operations."),
        ("message", "Sent to participants when a new message lands in a Communications thread (for inbound SMS from drivers/mechanics)."),
    ])

    add_heading(doc, "8.3. Configuring preferences", level=2, color=NAVY)
    add_para(doc, "Org-wide defaults: Admin → Settings → Notifications. Per-user overrides: each user's Profile → Notifications.")

    # ========================================================================
    # SECTION 9 — INTEGRATIONS
    # ========================================================================
    add_section_break(doc)
    add_heading(doc, "9. Integrations", level=1, color=AMBER)
    add_para(doc, "Five external systems are integrated. Each is configured via Supabase function secrets (set via 'supabase secrets set' on the operator's CLI), not via the UI.")

    add_heading(doc, "9.1. Twilio (SMS & Communications)", level=2, color=NAVY)
    add_para(doc, "Two distinct uses:")
    add_bullets(doc, [
        "Programmable SMS — sends one-way job-assignment notifications to drivers. Uses TWILIO_ACCOUNT_SID, TWILIO_AUTH_TOKEN, TWILIO_FROM_NUMBER.",
        "Conversations API — powers the Communications tab with two-way SMS threading, MMS attachments, delivery receipts, and webhook-based inbound. Uses TWILIO_API_KEY_SID, TWILIO_API_KEY_SECRET, TWILIO_CONVERSATIONS_SERVICE_SID, TWILIO_WEBHOOK_BASE_URL.",
    ])
    add_para(doc, "The twilio-verify edge function checks all six credentials and the live webhook URL configuration in the Twilio Console. Run it any time you suspect a misconfiguration.")

    add_heading(doc, "9.2. Geotab (GPS & telematics)", level=2, color=NAVY)
    add_para(doc, "Geotab provides vehicle GPS, odometer, engine hours, and movement data. Used by:")
    add_bullets(doc, [
        "Live vehicle map — every 5 minutes, the geotab-sync-locations cron polls and updates vehicle_locations",
        "Timesheet cross-reference — at clock-in time, the driver's GPS is compared to the vehicle's GPS at the same moment",
        "Prolonged stop detection — the prolonged-stop-check cron alerts admin when a vehicle stops for >2 hours mid-shift",
        "Preventive maintenance scheduling — odometer thresholds trigger MWO auto-creation",
    ])
    add_para(doc, "Configure via GEOTAB_USERNAME, GEOTAB_PASSWORD, GEOTAB_DATABASE secrets.")

    add_heading(doc, "9.3. QuickBooks Online (invoicing & payroll)", level=2, color=NAVY)
    add_para(doc, "Two push paths:")
    add_bullets(doc, [
        "qbo-push-invoice — when a work order is approved, generates a QBO Invoice using the client's rate table for line items",
        "qbo-push-time — exports the week's approved time entries as QBO TimeActivity records (for payroll)",
    ])
    add_para(doc, "Configure via QBO_CLIENT_ID, QBO_CLIENT_SECRET, QBO_REFRESH_TOKEN, QBO_REALM_ID, QBO_ENVIRONMENT (sandbox or production).")
    add_callout(doc, 'warning', "Until you're ready for production billing, leave QBO_ENVIRONMENT=sandbox. Set to production only after onboarding to your real QBO company.")

    add_heading(doc, "9.4. Fleetio (vehicle data import)", level=2, color=NAVY)
    add_para(doc, "One-time data migration from Fleetio. The fleetio-import edge function reads from Fleetio's API and bulk-imports vehicles, maintenance history, and fuel logs into the Yardward Pro tables. Configure via FLEETIO_BEARER_TOKEN secret. Run from /admin/vehicles → Import from Fleetio.")

    add_heading(doc, "9.5. Resend (tender digest emails)", level=2, color=NAVY)
    add_para(doc, "Resend is the email provider for the weekly tender digest. Configure via RESEND_API_KEY, TENDER_DIGEST_RECIPIENTS, TENDER_DIGEST_FROM secrets. If not configured, the tender-scrape function still runs but skips the email send.")
    add_callout(doc, 'tip', "Resend can also be configured as Supabase's SMTP provider for password reset emails (Studio → Auth → Settings → SMTP). This gives you reliable deliverability and your own branded sender domain.")

    # ========================================================================
    # SECTION 10 — SECURITY
    # ========================================================================
    add_section_break(doc)
    add_heading(doc, "10. Security & data isolation", level=1, color=AMBER)
    add_para(doc, "Yardward Pro is built with defense-in-depth security:")

    add_heading(doc, "10.1. Authentication", level=2, color=NAVY)
    add_bullets(doc, [
        "Supabase Auth handles all authentication — JWT-based with refresh tokens",
        "Passwords are bcrypt-hashed at rest",
        "Refresh tokens rotate on every use",
        "Failed login attempts trigger rate limiting",
        "Recovery emails are scoped one-time tokens — used or expired tokens cannot be replayed",
    ])

    add_heading(doc, "10.2. Authorization via Row-Level Security", level=2, color=NAVY)
    add_para(doc, "Every table in Yardward Pro has RLS enabled with policies that enforce role-based isolation. Examples:")
    add_field_table(doc, [
        ("Drivers", "Can SELECT only rows where driver_id = auth.uid()"),
        ("Mechanics", "Can SELECT only rows where mechanic_id = auth.uid() OR (role-based for shared shop resources)"),
        ("Admins", "Can SELECT all rows via is_admin() check"),
        ("INSERT policies", "WITH CHECK enforces user can only insert rows owned by themselves"),
        ("UPDATE policies", "Both USING and WITH CHECK prevent re-assignment of ownership"),
    ])
    add_callout(doc, 'info', "Even if a frontend bug or compromised client tried to query another user's data, Postgres would refuse to return the rows. RLS is enforced at the database engine level, below any application logic.")

    add_heading(doc, "10.3. Tokenized links (driver tokens)", level=2, color=NAVY)
    add_bullets(doc, [
        "256-bit random tokens via gen_random_bytes",
        "Stored hashed in the driver_tokens table (the raw token never appears in DB)",
        "Scope-locked — a token minted for /driver/tickets cannot access /driver/work-order",
        "One-shot — auto-burned on successful form submission",
        "Configurable expiry (default 24h)",
    ])

    add_heading(doc, "10.4. SECURITY DEFINER functions", level=2, color=NAVY)
    add_para(doc, "Sensitive operations (work order approval, PO approval, conversation creation, password rotation, etc.) run via SECURITY DEFINER Postgres functions. Each function:")
    add_bullets(doc, [
        "Has SET search_path = public, pg_temp to prevent search-path injection",
        "Includes an explicit role check (e.g. is_admin())",
        "Uses FOR UPDATE locks where atomicity matters (e.g. work order approval)",
        "Has REVOKE ALL FROM public, anon + GRANT EXECUTE TO authenticated, service_role",
    ])

    add_heading(doc, "10.5. Webhook signature verification", level=2, color=NAVY)
    add_para(doc, "Inbound webhooks (Twilio) are verified using HMAC-SHA1 signatures with constant-time byte comparison. Replay protection via the webhook_replay_log table prevents leaked-payload replay attacks. All input fields are allow-listed and regex-validated before any database write.")

    # ========================================================================
    # SECTION 11 — SUPPORT
    # ========================================================================
    add_section_break(doc)
    add_heading(doc, "11. Support & troubleshooting", level=1, color=AMBER)

    add_heading(doc, "11.1. For drivers and mechanics", level=2, color=NAVY)
    add_para(doc, "Submit support tickets through Profile → Help & support → Open a ticket. Your admin will get the ticket in the support_tickets table.")

    add_heading(doc, "11.2. For admins", level=2, color=NAVY)
    add_para(doc, "Common issues and where to investigate:")
    add_field_table(doc, [
        ("Driver can't sign in", "Settings → Users → find the driver → rotate password via the admin-rotate-password endpoint."),
        ("Driver isn't receiving SMS", "Check their phone number is real E.164 (Drivers tab → pencil icon). Check SMS log for delivery status. If 'failed', check Twilio Console for error code."),
        ("Form submission missing", "Check Error log → Dead-letter queue. May have failed retry budget. Click Requeue to retry."),
        ("Work order not generating invoice", "Check the client's rate table has matching line items for the work order's load_type. Check Error log for the qbo-push-invoice function."),
        ("Live map shows stale data", "Click Refresh now on the map page. If still stale, check Settings → Integrations → Geotab status. The geotab-sync-locations cron might be failing."),
        ("Communications messages not delivered as SMS", "Verify the recipient has a real E.164 phone. Check Error log for twilio-send-message. Run twilio-verify to confirm Conversations API configuration."),
    ])

    add_heading(doc, "11.3. For developers", level=2, color=NAVY)
    add_para(doc, "Yardward Pro is open-source. Issues, feature requests, and PRs welcome at https://github.com/ahsan-arch/yardward-pro.")

    # ========================================================================
    # SECTION 12 — GLOSSARY
    # ========================================================================
    add_section_break(doc)
    add_heading(doc, "12. Glossary", level=1, color=AMBER)
    add_field_table(doc, [
        ("Admin", "Org-level management role with full read/write access."),
        ("CVOR", "Commercial Vehicle Operator's Registration (Ontario). Regulatory framework for commercial vehicle operators."),
        ("Driver token", "Tokenized one-shot URL granting access to a specific form without signing in."),
        ("E.164", "International phone number format (e.g. +14165550100). Required by Twilio for SMS delivery."),
        ("Edge Function", "Server-side serverless function deployed to Supabase. Used for Twilio integration, QBO push, webhook handlers."),
        ("MWO", "Maintenance Work Order. Created when a vehicle needs repair or scheduled maintenance."),
        ("Operational", "Vehicle status when ready for use. Other statuses: maintenance (in shop), out_of_service (parked)."),
        ("PWA", "Progressive Web App. Installable from URL, works offline, no app store required."),
        ("PR / PO", "Purchase Request / Purchase Order. Submitted by mechanics for parts approval."),
        ("RLS", "Row-Level Security. Postgres feature that enforces row-by-row access control at the database engine."),
        ("RPC", "Remote Procedure Call. Postgres SECURITY DEFINER function callable from the client."),
        ("Service role", "Supabase API key with full database access, bypasses RLS. Used only by edge functions, never exposed to the client."),
        ("Supabase", "The backend-as-a-service platform Yardward Pro is built on (Postgres + Auth + Realtime + Storage + Edge Functions)."),
        ("Twilio Conversations API", "Twilio's multi-party threading API powering the Communications tab. Supports SMS + MMS + in-app webchat."),
        ("Vercel", "The hosting platform for the Yardward Pro frontend. Auto-deploys on push to main."),
        ("WO", "Work Order. Driver's on-site submission for a completed job. Once approved, generates an invoice."),
    ])

    # ------------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------------
    doc.save(DOCX)
    print(f"Saved: {DOCX}")


def convert_to_pdf():
    """Convert the .docx to PDF via Word COM automation."""
    print("Converting to PDF...")
    from docx2pdf import convert
    convert(str(DOCX), str(PDF))
    print(f"Saved: {PDF}")


if __name__ == "__main__":
    build()
    convert_to_pdf()
